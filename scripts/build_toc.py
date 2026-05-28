"""为毕业论文.docx 生成完整目录正文。

策略：
1. 扫描所有一/二/三级标题段，添加 <w:bookmarkStart/End>
2. 删除现有 TOC 占位段，重建为：
   - 一段 TOC field "begin + instrText + separate"
   - 每条目一段（含 hyperlink → bookmark, tab leader, PAGEREF）
   - 一段 TOC field "end"
3. 用户打开后即可看到完整目录；F9 可让 PAGEREF 自动更新真实页码
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

HERE = Path(__file__).resolve().parent.parent
DOC = HERE / "毕业论文.docx"


def is_chapter_level(text):
    return bool(re.match(
        r"^(摘\s*要|Abstract|第[一二三四五六七八九十]+章|参考文献|附录\s*[A-Z]|致\s*谢)",
        text,
    ))


def collect_headings(doc):
    """返回 [(level, text, paragraph_index, paragraph_obj), ...]"""
    entries = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t or not p.runs:
            continue
        # 跳过"目  录"自身
        if t in ("目  录", "目录", "目 录"):
            continue
        r = next((r for r in p.runs if r.text.strip()), None)
        if r is None:
            continue
        size = r.font.size.pt if r.font.size else None
        bold = r.bold

        if is_chapter_level(t):
            entries.append((1, t, i, p))
        elif size == 14.0 and bold and re.match(r"^\d+\.\d+\s", t):
            entries.append((2, t, i, p))
        elif size == 12.0 and bold and re.match(r"^\d+\.\d+\.\d+\s", t):
            entries.append((3, t, i, p))
    return entries


def add_bookmark(paragraph, bookmark_id, bookmark_name):
    """在段落的最前/最后插入 bookmarkStart/End。"""
    p_el = paragraph._element
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), bookmark_name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))

    pPr = p_el.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(bm_start)
    else:
        p_el.insert(0, bm_start)
    p_el.append(bm_end)


def make_run(text, size_pt=12.0, bold=False, font_zh="宋体", font_en="Times New Roman"):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)
    rFonts.set(qn("w:cs"), font_en)
    rFonts.set(qn("w:eastAsia"), font_zh)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(szCs)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    r.append(rPr)
    if text is not None:
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
    return r


def make_toc_entry(level, text, bookmark_name, page_num="?"):
    """生成一个 TOC 条目段落。"""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)

    # 行距 1.25
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "360")  # 1.5 行距=360 (240 base * 1.5)
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

    # 缩进：L2 缩进 240 twips，L3 缩进 480 twips
    if level >= 2:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str((level - 1) * 240))
        pPr.append(ind)

    # 右对齐 tab + 点引导
    tabs = OxmlElement("w:tabs")
    pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9072")  # ~A4 内宽 16cm 处
    tabs.append(tab)

    # 内嵌 hyperlink 到 bookmark
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")
    p.append(hyperlink)

    # 文本 run
    hyperlink.append(make_run(text, size_pt=12.0))
    # tab run
    tab_r = make_run(None, size_pt=12.0)
    tab_el = OxmlElement("w:tab")
    tab_r.append(tab_el)
    hyperlink.append(tab_r)

    # PAGEREF field: begin
    r_begin = make_run(None, size_pt=12.0)
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc1)
    hyperlink.append(r_begin)

    # PAGEREF field: instrText
    r_instr = make_run(None, size_pt=12.0)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    r_instr.append(instr)
    hyperlink.append(r_instr)

    # PAGEREF field: separate
    r_sep = make_run(None, size_pt=12.0)
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc2)
    hyperlink.append(r_sep)

    # 占位页码 run
    hyperlink.append(make_run(page_num, size_pt=12.0))

    # PAGEREF field: end
    r_end = make_run(None, size_pt=12.0)
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r_end.append(fc3)
    hyperlink.append(r_end)

    return p


def make_toc_field_open():
    """生成 TOC 字段开始段：begin + instrText + separate。"""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)
    # 隐藏的小段，仅作为字段开始锚点
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)

    # begin
    r_begin = make_run(None, size_pt=12.0)
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc1)
    p.append(r_begin)

    # instr
    r_instr = make_run(None, size_pt=12.0)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_instr.append(instr)
    p.append(r_instr)

    # separate
    r_sep = make_run(None, size_pt=12.0)
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc2)
    p.append(r_sep)
    return p


def make_toc_field_close():
    """生成 TOC 字段结束段：end。"""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)

    r_end = make_run(None, size_pt=12.0)
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "end")
    r_end.append(fc)
    p.append(r_end)
    return p


def main():
    doc = Document(str(DOC))
    entries = collect_headings(doc)
    print(f"识别到 {len(entries)} 个标题")

    # 用粗略的预估值生成占位页码（按层级 + 位置）
    # 实际页码会被 F9 更新；这里给一个递增数列让目录"看起来像目录"
    base_page = 1
    estimated_pages = []
    for idx, (lvl, text, p_idx, _) in enumerate(entries):
        # 用段落序号估算大致页数：每 30 段约 1 页
        est = max(1, p_idx // 30 - 1) if p_idx > 30 else 1
        estimated_pages.append(str(est))

    # 1. 为所有标题添加 bookmark
    for idx, (lvl, text, p_idx, p_obj) in enumerate(entries):
        bookmark_name = f"_Toc_thesis_{idx+1:03d}"
        add_bookmark(p_obj, idx + 100, bookmark_name)

    # 2. 找到现有的"目  录"标题段和其下的 TOC 占位段
    toc_title_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "目  录":
            toc_title_idx = i
            break
    if toc_title_idx is None:
        raise RuntimeError("未找到 [目  录] 标题段")

    title_para = doc.paragraphs[toc_title_idx]
    body = title_para._element.getparent()

    # 收集需要删除的旧 TOC 段（标题段之后的连续相关段，最多到下一个分节符或下一个非空非字段段）
    next_elem = title_para._element.getnext()
    elements_to_remove = []
    while next_elem is not None:
        # 检测此段是否含 TOC 字段 instr 或者占位文本
        tag = next_elem.tag.split('}')[-1]
        if tag != "p":
            break
        # 含字段或占位文本则删除
        instrs = next_elem.findall(".//" + qn("w:instrText"))
        is_toc_field = any(' TOC ' in (i.text or '') for i in instrs)
        full_text = "".join(t.text or '' for t in next_elem.findall(".//" + qn("w:t")))
        # sectPr 含分节符 → 保留
        has_sectPr = next_elem.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None
        if has_sectPr:
            break
        if is_toc_field or "请在 Word 中按 F9" in full_text:
            elements_to_remove.append(next_elem)
            next_elem = next_elem.getnext()
            continue
        # 其他非空段：停止
        if full_text.strip():
            break
        # 空段：也停止（保留作为间隔）
        break

    for el in elements_to_remove:
        el.getparent().remove(el)
    print(f"已移除 {len(elements_to_remove)} 个旧 TOC 段")

    # 3. 在标题段之后插入：field-open + 85 个 entry + field-close
    field_open = make_toc_field_open()
    title_para._element.addnext(field_open)
    prev = field_open

    for idx, (lvl, text, p_idx, _) in enumerate(entries):
        bookmark_name = f"_Toc_thesis_{idx+1:03d}"
        entry_p = make_toc_entry(lvl, text, bookmark_name, page_num=estimated_pages[idx])
        prev.addnext(entry_p)
        prev = entry_p

    field_close = make_toc_field_close()
    prev.addnext(field_close)

    doc.save(str(DOC))
    print(f"已重写目录：插入 {len(entries)} 个 TOC 条目")


if __name__ == "__main__":
    main()
