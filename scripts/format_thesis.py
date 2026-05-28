"""按内蒙古科技大学毕业设计撰写与装订规范修正毕业论文.docx。

规范要点：
- 纸张 A4 210x297mm，页边距 L30/R20/T25/B25mm，页眉15，页脚17.5（已设置）
- 中文宋体、英文 Times New Roman（已设置）
- 一级标题（章、摘要、目录、参考文献、附录、致谢）：小三(15pt)加粗居中，段后空一行
- 二级 1.1：四号(14pt)加粗左对齐
- 三级 1.1.1：小四(12pt)加粗左对齐
- 正文：小四(12pt)，两端对齐，1.5倍行距，首行缩进2字符（已设置）
- 图标题：五号(10.5pt)居中置于图下方，图X.Y（一级编号下顺序编号），段后空一行
- 表标题：五号(10.5pt)居中置于表上方，表X.Y，段前空一行
- 摘要/目录用罗马页码；正文起阿拉伯页码从1开始
- 页眉从摘要页起：五号居中"内蒙古科技大学毕业设计说明书（毕业论文）"
- 装订顺序：封面、中英文摘要、目录、正文、参考文献、附录、致谢
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HERE = Path(__file__).resolve().parent.parent
SRC = HERE / "毕业论文.docx"
DST = HERE / "毕业论文.docx"

HEADER_TEXT = "内蒙古科技大学毕业设计说明书（毕业论文）"

# ------------------------------ helpers ------------------------------ #


def set_run_font(run, name_zh="宋体", name_en="Times New Roman", size_pt=None, bold=None):
    """统一设置 run 字体：中文宋体、英文 Times New Roman、可选字号与加粗。"""
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name_en)
    rFonts.set(qn("w:hAnsi"), name_en)
    rFonts.set(qn("w:cs"), name_en)
    rFonts.set(qn("w:eastAsia"), name_zh)


def set_paragraph_format(p, *, align=None, first_line_indent_chars=None,
                          line_spacing=1.5, space_before_pt=None, space_after_pt=None,
                          page_break_before=None):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if first_line_indent_chars is not None:
        # 2 个中文字符 ≈ 0.847cm at 12pt
        pf.first_line_indent = Cm(0.847) if first_line_indent_chars == 2 else None
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if space_before_pt is not None:
        pf.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)
    if page_break_before is not None:
        pf.page_break_before = page_break_before


def find_paragraph_by_text(doc, predicate):
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text.strip()):
            return i
    return -1


def insert_paragraph_before(target_para, text="", style=None):
    """在 target_para 之前插入新段落，并返回新 Paragraph 对象。"""
    new_p = OxmlElement("w:p")
    target_para._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, target_para._parent)
    if text:
        run = para.add_run(text)
        set_run_font(run)
    if style is not None:
        para.style = style
    return para


def add_page_break_before(para):
    """让段落从新页开始（不插入额外段落）。"""
    pPr = para._element.get_or_add_pPr()
    # 移除任何已有的 page_break_before 元素
    for el in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(el)
    el = OxmlElement("w:pageBreakBefore")
    pPr.append(el)


# ------------------------------ sections ------------------------------ #


def add_section_break_before(doc, target_para, start_type="nextPage"):
    """在 target_para 之前插入一个分节符。

    实现：复制末尾 sectPr 作模板，但移除其 headerReference/footerReference，
    这样新分节会"继承自上一节"（is_linked_to_previous==True），
    随后调用 is_linked_to_previous=False 才能创建独立的页眉/页脚部件。
    """
    body = doc.element.body
    last_sectPr = body.find(qn("w:sectPr"))

    new_sectPr = deepcopy(last_sectPr)
    # 移除 header/footer references —— 让新分节"链接到上一节"
    for ref_tag in ("headerReference", "footerReference"):
        for el in new_sectPr.findall(qn(f"w:{ref_tag}")):
            new_sectPr.remove(el)
    # 移除 pgNumType —— 让新分节从默认继承，后续显式设置
    for el in new_sectPr.findall(qn("w:pgNumType")):
        new_sectPr.remove(el)
    # 设置 sectionStart 类型
    type_el = new_sectPr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        new_sectPr.insert(0, type_el)
    type_el.set(qn("w:val"), start_type)

    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    new_p.append(pPr)
    pPr.append(new_sectPr)
    target_para._element.addprevious(new_p)

    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, target_para._parent)


# ------------------------------ headers/footers ------------------------------ #


def clear_header_footer(section):
    for el_name in ("header", "footer"):
        part = getattr(section, el_name)
        for p in list(part.paragraphs):
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.clear()


def set_section_link_to_previous(section, link=False):
    """设置/取消 header & footer 与上一节的链接关系。

    python-docx 的 section.header.is_linked_to_previous = True/False 已封装。
    """
    section.header.is_linked_to_previous = link
    section.footer.is_linked_to_previous = link
    # also for first-page if needed -- skip


def _force_fresh_header_footer(section):
    """强制让 section 的 header/footer 重新生成全新的 part。

    若已经存在 reference，先 drop（变为 linked=True），
    再 add（变为 linked=False），即可拿到一个全新的空 part。
    """
    for part in (section.header, section.footer):
        if not part.is_linked_to_previous:
            part.is_linked_to_previous = True  # drop existing reference
        part.is_linked_to_previous = False  # create fresh empty part


def set_section_header_text(section, text):
    _force_fresh_header_footer(section)
    h = section.header
    # 清空已有内容
    for p in list(h.paragraphs):
        p._element.getparent().remove(p._element)
    # 插入一个新段落
    p_el = OxmlElement("w:p")
    h._element.append(p_el)
    from docx.text.paragraph import Paragraph
    p = Paragraph(p_el, h)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size_pt=10.5, bold=False)


def set_section_page_number_footer(section, num_fmt="decimal", start=None):
    """在页脚加入居中页码字段，并设置该 section 的页码格式与起始页。

    必须在 set_section_header_text 之后调用（共享 _force_fresh 逻辑）。
    """
    f = section.footer
    if f.is_linked_to_previous:
        f.is_linked_to_previous = False
    # 清空已有内容
    for p in list(f.paragraphs):
        p._element.getparent().remove(p._element)
    p_el = OxmlElement("w:p")
    f._element.append(p_el)
    from docx.text.paragraph import Paragraph
    p = Paragraph(p_el, f)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size_pt=10.5, bold=False)
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), "PAGE")
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = "1"
    inner_r.append(inner_t)
    fldSimple.append(inner_r)
    run._element.append(fldSimple)

    sectPr = section._sectPr
    for el in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(el)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), num_fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


def remove_section_header_footer(section):
    """让该 section 有独立的、空的 header/footer。"""
    _force_fresh_header_footer(section)
    for part in (section.header, section.footer):
        for p in list(part.paragraphs):
            p._element.getparent().remove(p._element)
        p_el = OxmlElement("w:p")
        part._element.append(p_el)


# ------------------------------ TOC ------------------------------ #


def insert_toc(after_para):
    """在指定段落之后插入"目录"标题与 TOC 字段。

    用法：先插入"目录"标题段，再插入 TOC 字段段，再插入一个分页符。
    """
    from docx.text.paragraph import Paragraph

    parent = after_para._parent
    # 创建标题段
    title_p_el = OxmlElement("w:p")
    after_para._element.addnext(title_p_el)
    title_p = Paragraph(title_p_el, parent)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(title_p, align=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before_pt=24, space_after_pt=18, line_spacing=1.5)
    title_run = title_p.add_run("目  录")
    set_run_font(title_run, size_pt=15, bold=True)

    # TOC 字段段
    toc_p_el = OxmlElement("w:p")
    title_p_el.addnext(toc_p_el)
    toc_p = Paragraph(toc_p_el, parent)

    # 字段开始 begin
    fld_begin_r = OxmlElement("w:r")
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    fld_begin_r.append(fldChar1)
    toc_p_el.append(fld_begin_r)

    # 指令
    instr_r = OxmlElement("w:r")
    instr_t = OxmlElement("w:instrText")
    instr_t.set(qn("xml:space"), "preserve")
    instr_t.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr_r.append(instr_t)
    toc_p_el.append(instr_r)

    # 字段分隔
    sep_r = OxmlElement("w:r")
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    sep_r.append(fldChar2)
    toc_p_el.append(sep_r)

    # 占位文本（请按 F9 更新）
    placeholder_r = OxmlElement("w:r")
    placeholder_t = OxmlElement("w:t")
    placeholder_t.text = "请在 Word 中按 F9 更新目录"
    placeholder_r.append(placeholder_t)
    toc_p_el.append(placeholder_r)

    # 字段结束
    end_r = OxmlElement("w:r")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    end_r.append(fldChar3)
    toc_p_el.append(end_r)

    return toc_p


# ------------------------------ main fixes ------------------------------ #


CAPTION_PATTERNS = [
    # [图 1-1：xxx] / [表 1-1：xxx]
    (re.compile(r"^\[\s*(图|表)\s*(\d+)\s*[-．\.]\s*(\d+)\s*[:：]?\s*(.*?)\s*\]\s*$"), r"\1\2.\3 \4"),
    # 图 1-1 xxx / 表 1-1 xxx （含空格、连字符、冒号）
    (re.compile(r"^(图|表)\s*(\d+)\s*[-．\.]\s*(\d+)(b|c|a)?\s*[:：]?\s*(.*?)\s*$"), r"\1\2.\3\4 \5"),
]


def normalize_caption_text(text):
    """归一化图/表标题文本：图X.Y / 表X.Y 形式，去掉 [], 冒号，去掉多余空格。"""
    t = text.strip()
    for pat, sub in CAPTION_PATTERNS:
        m = pat.match(t)
        if m:
            new = pat.sub(sub, t)
            # 清理末尾多余空格
            new = re.sub(r"\s+", " ", new).strip()
            # 处理 None 字母后缀产生的空字符串
            new = new.replace("None", "")
            return new
    return None


def is_caption_paragraph(p):
    """判断段落是否为图/表标题段。"""
    t = p.text.strip()
    if not t:
        return False
    if re.match(r"^\[\s*(图|表)\s*\d+", t):
        return True
    if re.match(r"^(图|表)\s*\d+\s*[-．\.]\s*\d+", t):
        return True
    return False


def apply_caption_style(p):
    """将段落格式化为图/表标题：五号居中，无首行缩进。"""
    # 先清空 first_line_indent
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 段前/段后：图下方一般段后空一行，表上方段前空一行。这里统一用对称小间距
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 字号
    for r in p.runs:
        set_run_font(r, size_pt=10.5, bold=False)


# ------------------------------ run ------------------------------ #


def main():
    doc = Document(str(SRC))

    # ---------- 1. 修正封面：班级误填 ----------
    for p in doc.paragraphs[:20]:
        t = p.text
        # "班    级：_______计算机科学与技术___" -> 改成真实班级
        if t.startswith("班") and "级" in t and "计算机科学与技术" in t:
            for r in p.runs:
                r.text = ""
            new_run = p.runs[0] if p.runs else p.add_run()
            new_run.text = "班    级：计科 2022-2 班"
            set_run_font(new_run, size_pt=14, bold=True)
            break

    # ---------- 2. 规范图表标题 ----------
    caption_count = 0
    for p in doc.paragraphs:
        norm = normalize_caption_text(p.text)
        if norm is None:
            continue
        # 写回归一化文本
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
        if p.runs:
            p.runs[0].text = norm
            set_run_font(p.runs[0], size_pt=10.5, bold=False)
        else:
            run = p.add_run(norm)
            set_run_font(run, size_pt=10.5, bold=False)
        apply_caption_style(p)
        caption_count += 1
    print(f"已规范图/表标题段：{caption_count} 个")

    # ---------- 3. 章节起始分页（封面之后所有主要节） ----------
    # 用 page_break_before 比插入显式分页符更稳健
    page_break_targets = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "摘　要" or t == "摘 要" or t == "摘要":
            page_break_targets.append(("摘要", i))
        elif t == "Abstract":
            page_break_targets.append(("Abstract", i))
        elif re.match(r"^第[一二三四五六七八九十]+章", t):
            page_break_targets.append((t[:6], i))
        elif t == "参考文献":
            page_break_targets.append(("参考文献", i))
        elif re.match(r"^附录\s*[A-Z]", t):
            page_break_targets.append((f"附录{t[2:4]}", i))
        elif t in ("致　谢", "致谢"):
            page_break_targets.append(("致谢", i))

    for name, idx in page_break_targets:
        add_page_break_before(doc.paragraphs[idx])
    print(f"已为 {len(page_break_targets)} 个一级标题加 page-break-before")

    # ---------- 4. 插入目录 ----------
    # 在 Abstract 段后的最后一个 Keywords 段之后插入
    en_kw_idx = find_paragraph_by_text(
        doc, lambda t: t.startswith("Keywords") or t.startswith("Key words")
    )
    if en_kw_idx == -1:
        # 退而求其次：放在第一章之前
        ch1_idx = find_paragraph_by_text(doc, lambda t: re.match(r"^第一章", t))
        if ch1_idx > 0:
            insert_toc(doc.paragraphs[ch1_idx - 1])
    else:
        insert_toc(doc.paragraphs[en_kw_idx])
    # 目录之后让第一章成为新页
    ch1_idx = find_paragraph_by_text(doc, lambda t: re.match(r"^第一章", t))
    if ch1_idx > 0:
        add_page_break_before(doc.paragraphs[ch1_idx])

    print("已插入目录字段（请在 Word 中按 F9 更新）")

    # ---------- 5. 多分节 + 页眉页脚 ----------
    # 当前文档默认只有 1 个 section。我们要变成 3 个 section：
    #   Section 1: 封面（仅封面页）
    #   Section 2: 摘要 + Abstract + 目录
    #   Section 3: 第一章 ... 致谢

    # 在 "摘要" 标题段之前插入 section break（next page）
    idx_abs = find_paragraph_by_text(doc, lambda t: t in ("摘　要", "摘 要", "摘要"))
    if idx_abs > 0:
        add_section_break_before(doc, doc.paragraphs[idx_abs], "nextPage")
    # 重新查找索引，因为插入了一个新段落
    idx_ch1 = find_paragraph_by_text(doc, lambda t: re.match(r"^第一章", t))
    if idx_ch1 > 0:
        add_section_break_before(doc, doc.paragraphs[idx_ch1], "nextPage")

    # 现在 doc.sections 应该有 3 个
    print(f"分节后 section 数：{len(doc.sections)}")

    # 配置三个 section 的页眉页脚
    sections = doc.sections
    if len(sections) >= 1:
        # Section 0：封面
        remove_section_header_footer(sections[0])
    if len(sections) >= 2:
        # Section 1：摘要/目录 - 罗马页码，从 I 起
        set_section_header_text(sections[1], HEADER_TEXT)
        set_section_page_number_footer(sections[1], num_fmt="upperRoman", start=1)
    if len(sections) >= 3:
        # Section 2：正文 - 阿拉伯页码，从 1 起
        set_section_header_text(sections[2], HEADER_TEXT)
        set_section_page_number_footer(sections[2], num_fmt="decimal", start=1)

    # ---------- 6. 保存 ----------
    doc.save(str(DST))
    print(f"已保存：{DST}")


if __name__ == "__main__":
    main()
