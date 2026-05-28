"""把 毕业论文.md 转换为符合内蒙古科技大学规范的 .docx

学校规范要点：
- 纸张 A4（210×297mm），页边距 上下 25mm / 左 30mm / 右 20mm
- 中文宋体、西文 Times New Roman
- 论文题目 / 一级标题（章）：小三号字加粗、居中
- 二级标题：四号字加粗、左对齐
- 三级标题：小四号字加粗、左对齐
- 正文：小四号字、两端对齐、1.5 倍行距、首行缩进 2 字符
- 图标题：五号字居中置于图下方
- 表标题：五号字居中置于表上方（脚本会自动把 markdown 表格后的"表X-Y"行提到表上方）
- 页眉：五号字居中"内蒙古科技大学毕业设计说明书（毕业论文）"
- 页脚：五号字居中页码
"""
import re
import sys
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm, Inches, Cm, RGBColor
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 让本脚本能从 scripts/ 目录直接运行（cwd 任意），导入同目录的公式模块
sys.path.insert(0, str(Path(__file__).resolve().parent))
try:
    from mathml_to_omml import latex_to_omath, latex_to_omath_para, split_tag  # noqa: E402
except ModuleNotFoundError:
    latex_to_omath = None
    latex_to_omath_para = None

    def split_tag(latex: str):
        m = re.search(r"\\tag\{([^}]+)\}", latex)
        if not m:
            return latex, None
        return re.sub(r"\\tag\{([^}]+)\}", "", latex).strip(), f"({m.group(1)})"


# ========== 学校规范常量 ==========
TITLE_PT = 18       # 封面"内蒙古科技大学" — 二号
SUBTITLE_PT = 16    # 封面副标题 — 三号
H1_PT = 15          # 一级标题（章）— 小三号
H2_PT = 14          # 二级标题（条）— 四号
H3_PT = 12          # 三级标题（款）— 小四号
H4_PT = 12          # 四级标题（项）— 小四号
BODY_PT = 12        # 正文 — 小四号
CAPTION_PT = 10.5   # 图表标题、页眉页脚 — 五号
CODE_PT = 10        # 代码块

CN_FONT = "宋体"
EN_FONT = "Times New Roman"
CODE_FONT = "Courier New"

PAGE_W_MM = 210
PAGE_H_MM = 297
MARGIN_TOP_MM = 25
MARGIN_BOTTOM_MM = 25
MARGIN_LEFT_MM = 30
MARGIN_RIGHT_MM = 20

HEADER_TEXT = "内蒙古科技大学毕业设计说明书（毕业论文）"
TOC_PAGES_JSON = "毕业论文.toc-pages.json"


def strip_markdown_marks(text: str) -> str:
    """去掉题注、封面行中遗留的基础 Markdown 标记。"""
    text = text.strip()
    text = re.sub(r"^\*\*(.*?)\*\*", r"\1", text)
    text = text.replace("**", "")
    return text.strip()


def normalize_caption_text(text: str) -> str:
    """把 Markdown/占位题注整理为学校常见的 图1.1 / 表1.1 形式。"""
    text = strip_markdown_marks(text)
    if text.startswith("[") and text.endswith("]"):
        text = text[1:-1].strip()
    text = text.replace("：", " ")
    text = re.sub(r"\s+", " ", text).strip()

    def repl(match):
        kind, chapter, item = match.groups()
        return f"{kind}{chapter}.{item}"

    text = re.sub(r"^(图|表)\s*(\d+)[-－](\d+[A-Za-z]?)", repl, text)
    text = re.sub(r"^(图|表)\s*(\d+)\.(\d+[A-Za-z]?)", repl, text)
    return text


def is_placeholder_caption(text: str) -> bool:
    s = text.strip()
    return bool(re.match(r"^\[(图|表)\s*\d+", s))


def is_table_caption(text: str) -> bool:
    s = strip_markdown_marks(text)
    return bool(re.match(r"^表\s*\d+", s))


def normalize_toc_key(text: str) -> str:
    """目录页码映射用的稳定 key，忽略空白和 PDF 抽取产生的空字符。"""
    text = strip_markdown_marks(text)
    return re.sub(r"[\s\u3000\x00]+", "", text)


# ========== 字体辅助 ==========
def set_run_font(run, size_pt=BODY_PT, bold=False, italic=False, code=False):
    """设置 run 字体：英文 Times New Roman，中文宋体（通过 eastAsia 属性）"""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

    rPr = run._element.get_or_add_rPr()
    # 清除已有 rFonts
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    if code:
        rFonts.set(qn("w:ascii"), CODE_FONT)
        rFonts.set(qn("w:hAnsi"), CODE_FONT)
        rFonts.set(qn("w:eastAsia"), CODE_FONT)
        rFonts.set(qn("w:cs"), CODE_FONT)
    else:
        rFonts.set(qn("w:ascii"), EN_FONT)
        rFonts.set(qn("w:hAnsi"), EN_FONT)
        rFonts.set(qn("w:eastAsia"), CN_FONT)
        rFonts.set(qn("w:cs"), EN_FONT)
    rPr.append(rFonts)


def set_style_font(style, size_pt=BODY_PT, bold=False, alignment=None):
    """设置 Word 样式字体、字号、段落节奏，保证目录可识别标题级别。"""
    style.font.name = EN_FONT
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    for old in rpr.findall(qn("w:rFonts")):
        rpr.remove(old)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), CN_FONT)
    rfonts.set(qn("w:cs"), EN_FONT)
    rpr.append(rfonts)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    if alignment is not None:
        style.paragraph_format.alignment = alignment


def setup_styles(doc):
    """建立学校规范所需的正文、标题、目录相关样式。"""
    normal = doc.styles["Normal"]
    set_style_font(normal, BODY_PT, bold=False)
    normal.paragraph_format.first_line_indent = Pt(BODY_PT * 2)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, H1_PT, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    h1.paragraph_format.first_line_indent = None
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(18)

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, H2_PT, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    h2.paragraph_format.first_line_indent = None
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, H3_PT, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    h3.paragraph_format.first_line_indent = None
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    h4 = doc.styles["Heading 4"]
    set_style_font(h4, H4_PT, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    h4.paragraph_format.first_line_indent = None
    h4.paragraph_format.space_before = Pt(6)
    h4.paragraph_format.space_after = Pt(3)

    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            st = doc.styles[style_name]
            set_style_font(st, BODY_PT, bold=False)
            st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            st.paragraph_format.line_spacing = 1.5


def add_runs_inline(p, text, size_pt=BODY_PT):
    """解析 **bold** *italic* `code` $latex$，逐段添加 run 或行内公式"""
    # 优先切出 $...$ 公式（避免和 *、` 冲突），再处理 markdown 标记
    # 正则：$ 后跟非 $ 字符直到下一个 $，要求成对，不允许跨行
    pattern = r"(\$[^$\n]+\$|\*\*[^*\n]+\*\*|\*[^*\n]+\*|`[^`\n]+`)"
    pieces = re.split(pattern, text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("$") and piece.endswith("$") and len(piece) >= 3:
            latex = piece[1:-1]
            try:
                if latex_to_omath is None:
                    raise RuntimeError("latex2mathml unavailable")
                omath = latex_to_omath(latex, display=False)
                p._element.append(omath)
            except Exception as e:
                print(f"[math-warning] 行内公式转换失败，回退为文本: {piece} ({e})", flush=True)
                r = p.add_run(piece)
                set_run_font(r, size_pt=size_pt)
        elif piece.startswith("**") and piece.endswith("**") and len(piece) >= 4:
            r = p.add_run(piece[2:-2])
            set_run_font(r, size_pt=size_pt, bold=True)
        elif piece.startswith("*") and piece.endswith("*") and len(piece) >= 3 and not piece.startswith("**"):
            r = p.add_run(piece[1:-1])
            set_run_font(r, size_pt=size_pt, italic=True)
        elif piece.startswith("`") and piece.endswith("`") and len(piece) >= 2:
            r = p.add_run(piece[1:-1])
            set_run_font(r, size_pt=size_pt, code=True)
        else:
            r = p.add_run(piece)
            set_run_font(r, size_pt=size_pt)


# ========== 段落构造 ==========
def add_body_paragraph(doc, text, indent=True, size_pt=BODY_PT):
    """正文段落：小四 / 1.5 倍行距 / 首行缩进 2 字符 / 两端对齐"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Pt(size_pt * 2)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    add_runs_inline(p, text, size_pt=size_pt)
    return p


def add_heading(doc, text, level):
    """添加标题。level=0 论文题目；1 章；2 节；3 款"""
    style_name = f"Heading {min(max(level, 1), 4)}"
    p = doc.add_paragraph(style=style_name)
    if level <= 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = H1_PT
        before, after = Pt(24), Pt(18)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = H2_PT
        before, after = Pt(12), Pt(6)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = H3_PT
        before, after = Pt(8), Pt(4)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = None
    pf.space_before = before
    pf.space_after = after
    r = p.add_run(text)
    set_run_font(r, size_pt=size, bold=True)
    return p


def add_cover_line(doc, text, size_pt=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=12):
    """封面行"""
    text = strip_markdown_marks(text)
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    add_runs_inline(p, text, size_pt=size_pt)
    if bold:
        for r in p.runs:
            r.bold = True


def add_caption(doc, text, position="below"):
    """图/表标题：五号字居中"""
    text = normalize_caption_text(text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = None
    if position == "below":
        pf.space_after = Pt(12)
        pf.space_before = Pt(0)
    else:
        pf.space_before = Pt(12)
        pf.space_after = Pt(0)
    add_runs_inline(p, text, size_pt=CAPTION_PT)
    return p


def add_image_block(doc, image_path, caption=None, max_width_in=5.5):
    """插入图片 + 标题"""
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_after = Pt(6)
    img_p.paragraph_format.first_line_indent = None
    run = img_p.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(max_width_in))
    except Exception as e:
        # 图片插入失败，输出占位
        p = doc.add_paragraph(f"[图片插入失败: {image_path} — {e}]")
    if caption:
        add_caption(doc, caption, position="below")


def add_display_math(doc, latex):
    """添加块级公式 $$...$$，居中显示，\\tag{X-Y} 编号通过 tab 右对齐"""
    cleaned, tag = split_tag(latex)
    try:
        if latex_to_omath is None:
            raise RuntimeError("latex2mathml unavailable")
        omath = latex_to_omath(cleaned, display=True)
    except Exception as e:
        print(f"[math-warning] 块公式转换失败，回退为文本: {latex} ({e})", flush=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = None
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        r = p.add_run(f"{cleaned} {tag or ''}".strip())
        set_run_font(r, size_pt=CAPTION_PT)
        return p

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # 居中公式 + 编号右对齐：用 tab stop。先公式，再 tab，再编号。
    # 行可用宽 = A4(210) - 左30 - 右20 = 160mm。设置中心 tab 在 80mm，右 tab 在 160mm。
    if tag:
        tab_stops = pf.tab_stops
        tab_stops.add_tab_stop(Mm(80), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Mm(160), WD_TAB_ALIGNMENT.RIGHT)
        # 公式前一个 tab → 推到中心；公式后一个 tab → 编号到右
        r_pre = p.add_run("\t")
        set_run_font(r_pre, size_pt=BODY_PT)
        p._element.append(omath)
        r_mid = p.add_run("\t")
        set_run_font(r_mid, size_pt=BODY_PT)
        r_tag = p.add_run(tag)
        set_run_font(r_tag, size_pt=BODY_PT)
        # 段落本身改为左对齐，让 tab stop 控制位置
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p._element.append(omath)
    return p


def add_md_table(doc, headers, rows, caption=None):
    """添加 markdown 表格转 docx 表"""
    if caption:
        add_caption(doc, caption, position="above")

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    cell_margins = OxmlElement("w:tblCellMar")
    for side, width in (("top", "80"), ("bottom", "80"), ("left", "120"), ("right", "120")):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), width)
        node.set(qn("w:type"), "dxa")
        cell_margins.append(node)
    tbl_pr.append(cell_margins)

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].text = ""
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        cp.paragraph_format.line_spacing = 1.2
        cp.paragraph_format.first_line_indent = None
        add_runs_inline(cp, h, size_pt=CAPTION_PT)
        for r in cp.runs:
            r.bold = True

    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            if j >= len(headers):
                continue
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].text = ""
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(val) <= 28 else WD_ALIGN_PARAGRAPH.LEFT
            cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            cp.paragraph_format.line_spacing = 1.2
            cp.paragraph_format.first_line_indent = None
            add_runs_inline(cp, val, size_pt=CAPTION_PT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = None


def add_code_block(doc, code_lines):
    """代码块：等宽字体，单倍行距"""
    for line in code_lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.left_indent = Cm(0.5)
        r = p.add_run(line if line else " ")
        set_run_font(r, size_pt=CODE_PT, code=True)


# ========== 页面/页眉/页脚/目录 ==========
def setup_section(section):
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    section.top_margin = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(17.5)


def clear_part(part):
    for p in list(part.paragraphs):
        p._element.getparent().remove(p._element)
    p_el = OxmlElement("w:p")
    part._element.append(p_el)
    return part.paragraphs[0]


def set_page_number_format(section, fmt="decimal", start=None):
    sect_pr = section._sectPr
    for el in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(el)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sect_pr.append(pg)


def setup_header_footer(section, header_text=HEADER_TEXT, page_fmt="decimal", start=None):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    header_p = clear_part(header)
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header_p.add_run(header_text)
    set_run_font(r, size_pt=CAPTION_PT)

    footer = section.footer
    footer_p = clear_part(footer)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run()
    set_run_font(run, size_pt=CAPTION_PT)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_page_number_format(section, fmt=page_fmt, start=start)


def remove_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_part(section.header)
    clear_part(section.footer)


def add_section(doc, start_type=WD_SECTION.NEW_PAGE):
    section = doc.add_section(start_type)
    setup_section(section)
    return section


def ensure_front_section(doc):
    if len(doc.sections) == 1:
        section = add_section(doc, WD_SECTION.NEW_PAGE)
        setup_header_footer(section, page_fmt="upperRoman", start=1)


def ensure_body_section(doc):
    section = add_section(doc, WD_SECTION.NEW_PAGE)
    setup_header_footer(section, page_fmt="decimal", start=1)
    return section


def to_roman(num: int) -> str:
    """把摘要/目录部分页码转为大写罗马数字。"""
    pairs = (
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    )
    out = []
    for value, symbol in pairs:
        while num >= value:
            out.append(symbol)
            num -= value
    return "".join(out)


def collect_toc_entries(md_path):
    """从 Markdown 标题收集三级目录；跳过封面和代码块里的 # 注释。"""
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    entries = []
    after_cover = False
    in_code = False
    for raw in lines:
        s = raw.strip()
        if s.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        if not after_cover:
            if s == "---":
                after_cover = True
            continue
        m = re.match(r"^(#{1,3})\s+(.+)$", s)
        if not m:
            continue
        title = strip_markdown_marks(m.group(2).strip())
        if title in ("目录", "目  录"):
            continue
        entries.append({"level": min(len(m.group(1)), 3), "title": title})
    return entries


def load_toc_page_map(md_path):
    """读取上一轮 PDF 反查得到的目录页码。没有文件时先生成空页码目录。"""
    json_path = Path(md_path).parent / TOC_PAGES_JSON
    if not json_path.exists():
        return {}
    try:
        raw = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception as exc:
        print(f"[toc-warning] 目录页码文件读取失败，将生成空页码目录: {exc}", flush=True)
        return {}
    return {normalize_toc_key(k): str(v) for k, v in raw.items()}


def add_toc(doc, entries=None, page_map=None):
    """插入可见三级目录。页码由 PDF 反查 JSON 二次生成写入。"""
    entries = entries or []
    page_map = page_map or {}

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("目  录")
    set_run_font(r, size_pt=H1_PT, bold=True)

    for entry in entries:
        title = entry["title"]
        level = int(entry["level"])
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = None
        pf.left_indent = Cm({1: 0, 2: 0.75, 3: 1.5}.get(level, 0))
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.tab_stops.add_tab_stop(Mm(160), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        r_title = p.add_run(title)
        set_run_font(r_title, size_pt=BODY_PT, bold=False)
        r_tab = p.add_run("\t")
        set_run_font(r_tab, size_pt=BODY_PT)
        r_page = p.add_run(page_map.get(normalize_toc_key(title), ""))
        set_run_font(r_page, size_pt=BODY_PT)


def enable_update_fields_on_open(doc):
    """让 Word 打开文件时更新目录/页码字段。"""
    settings = doc.settings.element
    for el in settings.findall(qn("w:updateFields")):
        settings.remove(el)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ========== Markdown 解析与渲染 ==========
def parse_md_table(lines, i):
    """返回 (headers, rows, next_i)。要求 lines[i] 是表头，lines[i+1] 是分隔行"""
    header_line = lines[i].strip().strip("|")
    headers = [c.strip() for c in header_line.split("|")]
    j = i + 2
    rows = []
    while j < len(lines) and lines[j].strip().startswith("|") and "|" in lines[j]:
        if "---" in lines[j]:
            j += 1
            continue
        row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
        rows.append(row)
        j += 1
    return headers, rows, j


def convert(md_path, docx_path):
    md_path = Path(md_path)
    docx_path = Path(docx_path)
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    section = doc.sections[0]
    setup_section(section)
    remove_header_footer(section)
    setup_styles(doc)

    i = 0
    in_code = False
    code_lines = []
    cover_done = False
    toc_done = False
    body_section_done = False
    last_progress = -1

    while i < len(lines):
        # 进度日志（每 100 行）
        if i // 100 != last_progress:
            last_progress = i // 100
            print(f"[progress] line {i}/{len(lines)}", flush=True)
        raw = lines[i]
        line = raw.rstrip()
        s = line.strip()

        # 代码块
        if s.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # 处理封面（前两行连续 # 标题 + 后续封面信息）
        if not cover_done and i == 0 and s.startswith("# "):
            add_cover_line(doc, s[2:].strip(), size_pt=TITLE_PT, bold=True)
            i += 1
            # 紧跟另一个 #
            while i < len(lines) and lines[i].strip().startswith("# "):
                add_cover_line(doc, lines[i].strip()[2:].strip(), size_pt=SUBTITLE_PT, bold=True)
                i += 1
            # 留几行空行
            for _ in range(3):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
            # 处理封面信息行（**xxx**: yyy）直到 ---
            while i < len(lines):
                cs = lines[i].strip()
                if cs == "---":
                    cover_done = True
                    i += 1
                    # 封面后插分页
                    doc.add_page_break()
                    break
                if not cs:
                    i += 1
                    continue
                add_cover_line(doc, cs, size_pt=14, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=8)
                i += 1
            continue

        # 一级标题（章 / 摘要 / 参考文献 / 附录 / 致谢）
        if s.startswith("# "):
            title = s[2:].strip()
            if title in ("摘　要", "摘 要", "摘要"):
                ensure_front_section(doc)
            elif title == "Abstract":
                doc.add_page_break()
            elif title.startswith("第一章"):
                if not toc_done:
                    doc.add_page_break()
                    add_toc(doc)
                    toc_done = True
                if not body_section_done:
                    body_section_done = True
                    ensure_body_section(doc)
            elif any(title.startswith(p) for p in ("第", "参考文献", "附录", "致　谢", "致谢")):
                doc.add_page_break()
            add_heading(doc, title, 1)
            i += 1
            continue

        # 二级标题
        if s.startswith("## "):
            add_heading(doc, s[3:].strip(), 2)
            i += 1
            continue

        # 三级标题
        if s.startswith("### "):
            add_heading(doc, s[4:].strip(), 3)
            i += 1
            continue

        # 四级标题
        if s.startswith("#### "):
            add_heading(doc, s[5:].strip(), 4)
            i += 1
            continue

        # 块公式 $$...$$ — 必须先于图片/表格/普通段落识别
        # 支持单行：$$latex$$ 和多行：$$\nlatex\n$$
        if s.startswith("$$"):
            inner = s[2:]
            # 单行结束
            if inner.endswith("$$") and len(inner) >= 2:
                latex = inner[:-2].strip()
                add_display_math(doc, latex)
                i += 1
                continue
            # 多行：收集直到下一个 $$
            j = i + 1
            buf = [inner]
            while j < len(lines):
                cur = lines[j].rstrip()
                if cur.strip().endswith("$$"):
                    buf.append(cur.strip()[:-2])
                    j += 1
                    break
                buf.append(cur)
                j += 1
            latex = "\n".join(b for b in buf if b is not None).strip()
            add_display_math(doc, latex)
            i = j
            continue

        # 图片
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", s)
        if m:
            alt, img_rel = m.group(1), m.group(2)
            full_path = (md_path.parent / img_rel).resolve()
            caption = None
            # 后续若有 **图 X-Y ... 行，作为图标题
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line.startswith("**图") or next_line.startswith("**表"):
                    caption = next_line
                    i += 1
            if full_path.exists():
                add_image_block(doc, full_path, caption=caption)
            else:
                add_body_paragraph(doc, f"[图片缺失: {img_rel}]", indent=False)
            i += 1
            continue

        # 表格
        if s.startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1] and "|" in lines[i + 1]:
            headers, rows, next_i = parse_md_table(lines, i)
            # 查找表标题（位于表后第一段非空文本，形如 "表X-Y ..." 或 "**表X-Y** ..."）
            caption = None
            k = next_i
            while k < len(lines) and not lines[k].strip():
                k += 1
            if k < len(lines):
                cand = lines[k].strip()
                m2 = re.match(r"^表\s*\d+", cand)
                m3 = re.match(r"^\*\*表\s*\d+", cand)
                if m2 or m3:
                    caption = cand
                    next_i = k + 1
            add_md_table(doc, headers, rows, caption=caption)
            i = next_i
            continue

        # 占位图题/表题行。图题单独成题注；表题通常已由表格逻辑前置处理。
        if is_placeholder_caption(s):
            add_caption(doc, s, position="below")
            i += 1
            continue
        if is_table_caption(s):
            add_caption(doc, s, position="above")
            i += 1
            continue

        # 列表项
        if s.startswith("- ") or s.startswith("* "):
            content = s[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.5
            add_runs_inline(p, content, size_pt=BODY_PT)
            i += 1
            continue

        # 有序列表
        m = re.match(r"^(\d+)\.\s+(.+)$", s)
        if m:
            content = m.group(2)
            p = doc.add_paragraph(style="List Number")
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.5
            add_runs_inline(p, content, size_pt=BODY_PT)
            i += 1
            continue

        # 分隔线
        if s == "---":
            i += 1
            continue

        # 空行
        if not s:
            i += 1
            continue

        # 普通段落
        add_body_paragraph(doc, line)
        i += 1

    enable_update_fields_on_open(doc)
    doc.save(docx_path)
    print(f"已保存: {docx_path}")
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")


if __name__ == "__main__":
    md = "/Users/a0000/Desktop/学业项目/毕业设计/毕业论文.md"
    docx = "/Users/a0000/Desktop/学业项目/毕业设计/毕业论文.docx"
    convert(md, docx)
